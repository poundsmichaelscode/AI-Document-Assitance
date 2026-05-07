from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from pypdf import PdfReader
class DocumentParserService:
    def parse(self, file_path: str, file_type: str, file_name: str):
        if file_type == 'pdf':
            reader = PdfReader(file_path); docs = []
            for index, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                if text.strip(): docs.append(LCDocument(page_content=text, metadata={'file_name': file_name, 'page': index + 1}))
            return docs
        if file_type == 'docx':
            doc = DocxDocument(file_path)
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            return [LCDocument(page_content=text, metadata={'file_name': file_name, 'page': 1})]
        if file_type == 'txt':
            text = Path(file_path).read_text(encoding='utf-8', errors='ignore')
            return [LCDocument(page_content=text, metadata={'file_name': file_name, 'page': 1})]
        raise ValueError('Unsupported file type')
